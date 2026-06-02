"""
agent_utils.py — Shared utilities for notes_agent.py, quiz_agent.py, enrich_agent.py
"""

import os
from pathlib import Path
from docx import Document


# ─── Load config.env automatically ───────────────────────────────────────────

def _load_config_env():
    """Read config.env from the script's directory and set missing env vars."""
    config_path = Path(__file__).parent / 'config.env'
    if not config_path.exists():
        return
    for line in config_path.read_text(encoding='utf-8').splitlines():
        line = line.strip()
        if not line or line.startswith('#'):
            continue
        if '=' in line:
            key, _, value = line.partition('=')
            key = key.strip()
            value = value.strip()
            if key and value and 'PASTE-YOUR-KEY' not in value:
                os.environ[key] = value

_load_config_env()

# ─── Topic registry ──────────────────────────────────────────────────────────

DOCX_MAP = {
    'T01': 'Oral Material/ETO_v2_Topic1_Alternator.docx',
    'T02': 'Oral Material/ETO_v2_Topic2_HV.docx',
    'T03': 'Oral Material/ETO_v2_Topic3_Motors.docx',
    'T04': 'Oral Material/ETO_v2_Topic4_Switchboard.docx',
    'T05': 'Oral Material/ETO_v2_Topic5_Transformers.docx',
    'T06': 'Oral Material/ETO_v2_Topic6_Sensors.docx',
    'T07': 'Oral Material/ETO_v2_Topic7_Control_PLC.docx',
    'T08': 'Oral Material/ETO_v2_Topic8_Electronics.docx',
    'T09': 'Oral Material/ETO_v2_Topic9_ShipMachinery_v2.docx',
    'T10': 'Oral Material/ETO_v2_Topic10_ICCP.docx',
    'T11': 'Oral Material/ETO_v2_Topic11_PowerFactor.docx',
    'T12': 'Oral Material/ETO_v2_Topic12_Cables.docx',
    'T13': 'Oral Material/ETO_v2_Topic13_BridgeEquip1.docx',
    'T14': 'Oral Material/ETO_v2_Topic14_BridgeEquip2.docx',
    'T15': 'Oral Material/ETO_v2_Topic15_FireFighting.docx',
    'T16': 'Oral Material/ETO_v2_Topic16_SOLAS.docx',
    'T17': 'Oral Material/ETO_v2_Topic17_MARPOL.docx',
    'T18': 'Oral Material/ETO_v2_Topic18_ElecSurvey.docx',
    'T19': 'Oral Material/ETO_v2_Topic19_Tanker.docx',
    'T20': 'Oral Material/ETO_v2_Topic20_Construction.docx',
    'T21': 'Oral Material/ETO_v2_Topic21_LSA.docx',
    'T22': 'Oral Material/ETO_v2_Topic22_Practical.docx',
    'T23': 'Oral Material/ETO_v2_Topic23_BatteriesDC.docx',
}

TOPIC_NAMES = {
    'T01': 'Alternator & Generator',
    'T02': 'High Voltage Systems',
    'T03': 'Electric Motors & Starters',
    'T04': 'Switchboard & Circuit Breakers',
    'T05': 'Transformers',
    'T06': 'Sensors & Instrumentation',
    'T07': 'Control Systems & PLC',
    'T08': 'Electronics',
    'T09': 'Ship Machinery',
    'T10': 'ICCP & Corrosion Protection',
    'T11': 'Power Factor & Harmonics',
    'T12': 'Cables & Wiring',
    'T13': 'Bridge Equipment I',
    'T14': 'Bridge Equipment II',
    'T15': 'Fire Fighting Systems',
    'T16': 'SOLAS',
    'T17': 'MARPOL',
    'T18': 'Electrical Survey',
    'T19': 'Tanker Electrics',
    'T20': 'Ship Construction',
    'T21': 'LSA & Safety Equipment',
    'T22': 'Practical Skills',
    'T23': 'Batteries & DC Systems',
}

# Quiz topic key (used in window.loadQuizzes() call and in question IDs)
TOPIC_KEYS = {
    'T01': 'T01_Alternator',
    'T02': 'T02_HV',
    'T03': 'T03_Motors',
    'T04': 'T04_Switchboard',
    'T05': 'T05_Transformers',
    'T06': 'T06_Sensors',
    'T07': 'T07_Control_PLC',
    'T08': 'T08_Electronics',
    'T09': 'T09_ShipMachinery',
    'T10': 'T10_ICCP',
    'T11': 'T11_PowerFactor',
    'T12': 'T12_Cables',
    'T13': 'T13_BridgeEquip1',
    'T14': 'T14_BridgeEquip2',
    'T15': 'T15_FireFighting',
    'T16': 'T16_SOLAS',
    'T17': 'T17_MARPOL',
    'T18': 'T18_ElecSurvey',
    'T19': 'T19_Tanker',
    'T20': 'T20_Construction',
    'T21': 'T21_LSA',
    'T22': 'T22_Practical',
    'T23': 'T23_BatteriesDC',
}

# ─── Docx extraction ─────────────────────────────────────────────────────────

def find_docx(topic_id):
    """Find the docx file for a topic, checking several relative paths."""
    rel = DOCX_MAP.get(topic_id)
    if not rel:
        return None
    candidates = [
        Path(rel),
        Path('..') / Path(rel).name,
        Path(rel.replace('Oral Material/', '')),
    ]
    for p in candidates:
        if p.exists():
            return p
    return None


def _para_font_size(para):
    """Return the most common run font size (in pt) for a paragraph, or None."""
    sizes = []
    for run in para.runs:
        if run.font and run.font.size:
            sizes.append(run.font.size.pt)
    if not sizes:
        return None
    # Most common size
    return max(set(sizes), key=sizes.count)


def _para_is_bold(para):
    """Return True if any non-whitespace run in this paragraph is bold."""
    for run in para.runs:
        if run.text.strip() and run.bold:
            return True
    return False


def extract_docx(docx_path):
    """
    Extract ALL content from a docx file: paragraphs and tables, in document order.
    Returns a list of dicts:
        {'type': 'para', 'text': str, 'heading': bool, 'font_size': float|None}
        {'type': 'table', 'rows': list[list[str]]}

    heading=True marks paragraphs that look like section/sub-section headers
    (bold + font_size >= 12pt, or Word Heading style).
    """
    doc = Document(str(docx_path))

    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    content = []
    for child in doc.element.body:
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            para = para_map.get(id(child))
            if para:
                text = para.text.strip()
                if not text:
                    continue

                style_name = (para.style.name or '') if para.style else ''
                is_word_heading = style_name.startswith('Heading')
                font_size = _para_font_size(para)
                is_bold = _para_is_bold(para)

                # Mark as heading: Word Heading style, OR bold + sizeable font
                heading = (
                    is_word_heading
                    or (is_bold and font_size is not None and font_size >= 12.0)
                )

                content.append({
                    'type': 'para',
                    'text': text,
                    'heading': heading,
                    'font_size': font_size,
                })

        elif tag == 'tbl':
            table = table_map.get(id(child))
            if table:
                rows = []
                for row in table.rows:
                    seen_ids = set()
                    cells = []
                    for cell in row.cells:
                        cell_id = id(cell._element)
                        if cell_id not in seen_ids:
                            cells.append(cell.text.strip())
                            seen_ids.add(cell_id)
                    if any(cells):
                        rows.append(cells)
                if rows:
                    content.append({'type': 'table', 'rows': rows})

    return content


def content_to_text(content):
    """
    Render extracted content as plain-text for feeding into prompts.
    Tables are rendered as pipe-delimited rows wrapped in [TABLE]...[/TABLE].
    """
    lines = []
    for item in content:
        if item['type'] == 'para':
            lines.append(item['text'])
        elif item['type'] == 'table':
            lines.append('[TABLE]')
            for row in item['rows']:
                lines.append(' | '.join(row))
            lines.append('[/TABLE]')
    return '\n'.join(lines)


# ─── Section detection ───────────────────────────────────────────────────────

import re

_NUMBERED_HEADING = re.compile(
    r'^(\d{1,2})[\.:\)]\s+(.{4,80})$'
)

# Titles to skip — they are document-level labels, not section headers
_SKIP_TITLES = {
    'ETO MMD ORAL EXAMINATION',
    'COMPREHENSIVE STUDY NOTES',
    'COLOUR CODE LEGEND',
    'MOST ASKED',
}


def detect_sections(content):
    """
    Find major section headers in extracted content.

    Strategy (in order of priority):
    1. Bold/heading-styled paragraphs (font_size 12-15pt, or Word Heading style).
       These are subsection/section titles as authored in the docx.
    2. ALL-CAPS lines as final fallback.
    3. Numbered headings where the title is ALL-CAPS and >= 12 chars
       (strict check to avoid matching numbered list items).

    Skips document-level labels (font_size >= 16pt) and known noise strings.

    Returns list of (section_num: int, title: str, start_index: int).
    """
    # Pass 1 — bold/heading paragraphs in the 12-15pt range
    styled = []
    for i, item in enumerate(content):
        if item['type'] != 'para':
            continue
        if not item.get('heading'):
            continue
        fs = item.get('font_size') or 0
        if fs >= 16:       # document title / subtitle — skip
            continue
        text = item['text'].strip()
        if len(text) < 5 or text.startswith('─') or text.startswith('—'):
            continue
        upper = text.upper()
        if any(skip in upper for skip in _SKIP_TITLES):
            continue
        styled.append((len(styled) + 1, text, i))

    if styled:
        return styled

    # Pass 2 — ALL-CAPS paragraphs (≥10 chars, not separators)
    caps = []
    for i, item in enumerate(content):
        if item['type'] != 'para':
            continue
        text = item['text'].strip()
        if (text == text.upper()
                and 10 < len(text) < 90
                and not text.replace(' ', '').replace('─', '').replace('—', '').isdigit()):
            upper = text.upper()
            if any(skip in upper for skip in _SKIP_TITLES):
                continue
            if text.startswith('─') or text.startswith('—'):
                continue
            caps.append((len(caps) + 1, text, i))

    if caps:
        return caps

    # Pass 3 — strict numbered headings: "N. TITLE IN CAPS (>= 12 chars)"
    _STRICT_NUM = re.compile(r'^(\d{1,2})[\.:\)]\s+([A-Z][A-Z &\-/()]{10,70})$')
    numbered = []
    for i, item in enumerate(content):
        if item['type'] != 'para':
            continue
        text = item['text'].strip()
        m = _STRICT_NUM.match(text)
        if m:
            numbered.append((int(m.group(1)), m.group(2).strip(), i))

    return numbered


def split_by_sections(content, sections):
    """
    Split content list into per-section slices.
    Returns list of (num_str: str, title: str, content_slice: list).

    If there is substantial content before the first section header, it is
    returned as a leading 'Introduction / Overview' section.
    If no sections found, returns a single ('1', 'All Content', content) entry.
    """
    if not sections:
        return [('1', 'All Content', content)]

    result = []

    # Preamble: content before the first detected section
    first_start = sections[0][2]
    if first_start > 0:
        preamble = content[:first_start]
        preamble_text = content_to_text(preamble)
        if len(preamble_text.strip()) > 200:   # only include if substantial
            result.append(('0', 'Introduction & Overview', preamble))

    for idx, (num, title, start) in enumerate(sections):
        end = sections[idx + 1][2] if idx + 1 < len(sections) else len(content)
        result.append((str(num), title, content[start:end]))

    return result


def make_cat_code(title):
    """Generate a 2-5 letter category abbreviation from a section title."""
    stop = {'AND', 'THE', 'OF', 'IN', 'FOR', 'TO', 'A', 'AN', 'VS', 'OR',
            'ITS', 'WITH', 'ON', 'BY', 'AT', 'AS', 'FROM'}
    words = [w for w in re.sub(r'[^A-Z ]', '', title.upper()).split() if w not in stop]
    if not words:
        return title[:4].upper()
    if len(words) == 1:
        return words[0][:5]
    return ''.join(w[0] for w in words[:5])
