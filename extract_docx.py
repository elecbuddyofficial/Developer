import docx
import sys

def extract_text(filepath):
    try:
        doc = docx.Document(filepath)
        text = []
        for p in doc.paragraphs:
            if p.text.strip():
                text.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip().replace('\n', ' '))
                if row_data:
                    text.append(" | ".join(row_data))
        with open('extracted_text.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(text))
        print(f"Extracted {len(text)} blocks to extracted_text.txt")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    extract_text(sys.argv[1])
